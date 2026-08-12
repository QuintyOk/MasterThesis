"""Describe the final SCE sample and selected supporting outcomes.

The script reports the sample composition, current modal split, commute and
availability measures, vanpool familiarity, experiment allocation, parking-
rights responses, and commute-distance bands. It also creates the modal-split
figure and the internal open-answer document used during the thesis process.

The formatted employee-level data required by this script are confidential."""

import os
import re
from difflib import get_close_matches

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document

# Load the task-level SCE dataset.
data_formatted = pd.read_excel(
    r"data\Access & Mobility Survey Formatted Data.xlsx",
    engine="openpyxl",
)

print("\n" + "=" * 80)

print("Columns data formatted: ", data_formatted.columns)

print("The number of respondents after data cleaning: ", data_formatted["ResponseId"].nunique())

print("\n" + "=" * 80)

print(data_formatted["primary_transport_mode"].value_counts())

# Calculate the respondent-level primary-mode distribution.
modal_split = (
    data_formatted[["ResponseId", "primary_transport_mode"]]
    .drop_duplicates("ResponseId")["primary_transport_mode"]
    .value_counts(dropna=False)
)

print("Current modal split (based on primary transport mode):")
print(modal_split)

print("\n" + "=" * 80)

# Aggregate the detailed modes to the three reporting categories.
mode_mapping = {
    'car_asml': 'Car',
    'bike': 'Bike',
    'moped': 'Bike',
    'carpool': 'Public Transport',
    'train_ebike': 'Public Transport',
    'train_bus': 'Public Transport',
    'vanpool': 'Public Transport',
    'car_pr': 'Public Transport',
    'bus': 'Public Transport'
}

modal_split = (
    data_formatted[['ResponseId', 'primary_transport_mode']]
    .drop_duplicates('ResponseId')
)

modal_split['category'] = modal_split['primary_transport_mode'].map(mode_mapping)

category_counts = modal_split['category'].value_counts()

print("# Modal split (3 categories):")
print(category_counts)

total = category_counts.sum()
labels = [
    f"{cat} ({count/total*100:.1f}%)"
    for cat, count in category_counts.items()
]

colors = ['#003A8F', '#00A3E0', '#A7A8AA']

# Plot the three-category modal split used during analysis.
plt.figure()
plt.pie(
    category_counts,
    labels=labels,
    colors=colors,
    startangle=90
)

plt.title("Modal Split (3 Categories)")

output_path = "plots/modal_split_pie.png"
os.makedirs("plots", exist_ok=True)
plt.savefig(output_path, bbox_inches='tight')

print(f"# Pie chart saved to: {output_path}")

plt.close()

print("\n" + "=" * 80)

# Summarise commute distance and office attendance.
commute_stats = (
    data_formatted[["ResponseId", "commute_distance_km"]]
    .drop_duplicates("ResponseId")["commute_distance_km"]
    .describe()
)

print("Descriptive statistics for commute distance (km):")
print(commute_stats)

print("\n" + "=" * 80)

office_days_stats = (
    data_formatted[["ResponseId", "office_days_per_week"]]
    .drop_duplicates("ResponseId")["office_days_per_week"]
    .describe()
)

print("Descriptive statistics for office days per week:")
print(office_days_stats)

print("\n" + "=" * 80)

# Summarise respondent-level mode availability.
availability_cols = ["car_available", "bike_available", "vanpool_available"]

pd.set_option("display.max_rows", None)
pd.set_option("display.max_columns", None)
pd.set_option("display.width", None)
pd.set_option("display.max_colwidth", None)

availability_cols = ["car_available", "bike_available", "vanpool_available"]


def yes_no(series):
    """Collapse detailed availability responses to Yes and No."""
    return series.apply(lambda x: "Yes" if isinstance(x, str) and x.startswith("Yes") else "No")

availability_yesno_pct = (
    data_formatted[["ResponseId"] + availability_cols]
    .drop_duplicates("ResponseId")[availability_cols]
    .apply(yes_no)
    .apply(lambda s: s.value_counts(normalize=True) * 100)
    .round(1)
)

print("Percentage availability (Yes/No, respondent level):")
print(availability_yesno_pct)

print("\n" + "=" * 80)

# Summarise the respondent-specific travel-time estimates.
time_cols = ["car_time", "pr_time", "bike_time", "pt_time", "vp_time"]

time_stats = (
    data_formatted[["ResponseId"] + time_cols]
    .drop_duplicates("ResponseId")[time_cols]
    .describe()
    .round(1)
)

print("Descriptive statistics for commute times (minutes, NaNs excluded):")
print(time_stats)

print("\n" + "=" * 80)

# Describe familiarity with and feasibility of vanpool.
familiar_vanpool_pct = (
    data_formatted[["ResponseId", "familiar_vanpool"]]
    .drop_duplicates("ResponseId")["familiar_vanpool"]
    .dropna()
    .value_counts(normalize=True)
    .mul(100)
    .round(1)
)

print("Familiarity with vanpool (in %, NaNs excluded):")
print(familiar_vanpool_pct)

print("\n" + "=" * 80)

vanpool_loc_counts = (
    data_formatted[["ResponseId", "vanpool_convenient_existing_location"]]
    .drop_duplicates("ResponseId")["vanpool_convenient_existing_location"]
    .dropna()
    .value_counts()
)

vanpool_loc_pct = (vanpool_loc_counts / vanpool_loc_counts.sum() * 100).round(1)

vanpool_loc_table = pd.DataFrame({
    "count": vanpool_loc_counts,
    "percentage": vanpool_loc_pct
})

print("Convenient existing vanpool location (counts and percentages, NaNs excluded):")
print(vanpool_loc_table)

print("\n" + "=" * 80)

vanpool_like_counts = (
    data_formatted[["ResponseId", "vanpool_likelihood_if_available"]]
    .drop_duplicates("ResponseId")["vanpool_likelihood_if_available"]
    .dropna()
    .value_counts()
)

vanpool_like_pct = (vanpool_like_counts / vanpool_like_counts.sum() * 100).round(1)

vanpool_like_table = pd.DataFrame({
    "count": vanpool_like_counts,
    "percentage": vanpool_like_pct
})

print("Likelihood of choosing vanpool if available (counts and percentages, NaNs excluded):")
print(vanpool_like_table)

print("\n" + "=" * 80)


def normalize_city(x):
    """Normalize an open-text city name before grouping variants."""
    if pd.isna(x):
        return pd.NA
    x = x.lower().strip()
    x = re.sub(r"[^\w\s]", "", x)
    x = re.sub(r"\s+", " ", x)
    return x

# Clean and group desired vanpool locations entered as free text.
vanpool_city_clean = (
    data_formatted[["ResponseId", "vanpool_desired_city"]]
    .drop_duplicates("ResponseId")
    .assign(city_clean=lambda d: d["vanpool_desired_city"].apply(normalize_city))
)


def cluster_strings(strings, cutoff=0.85):
    """Group similar city-name strings using fuzzy matching."""
    clusters = {}
    for s in strings:
        if s is pd.NA:
            continue
        match = next((k for k in clusters if get_close_matches(s, [k], n=1, cutoff=cutoff)), None)
        if match:
            clusters[match].append(s)
        else:
            clusters[s] = [s]
    return clusters

clusters = cluster_strings(vanpool_city_clean["city_clean"].dropna().unique())

city_map = {v: k for k, vs in clusters.items() for v in vs}
vanpool_city_clean["city_grouped"] = vanpool_city_clean["city_clean"].map(city_map)

city_counts = (
    vanpool_city_clean["city_grouped"]
    .dropna()
    .value_counts()
)

city_pct = (city_counts / city_counts.sum() * 100).round(1)

city_table = pd.DataFrame({
    "count": city_counts,
    "percentage": city_pct
})

print("Desired vanpool city (counts and percentages, cleaned):")
print(city_table)

print("\n" + "=" * 80)

# Describe allocation across the six availability experiments.
experiment_labels = {
    1: "Vanpool, PT",
    2: "Bike, Vanpool, PT",
    3: "Car, P+R, PT",
    4: "Bike, Car, P+R, PT",
    5: "Vanpool, Car, P+R, PT",
    6: "Bike, Vanpool, Car, P+R, PT",
}

experiment_counts = (
    data_formatted[["ResponseId", "experiment"]]
    .drop_duplicates("ResponseId")["experiment"]
    .value_counts(dropna=False)
)

experiment_pct = (experiment_counts / experiment_counts.sum() * 100).round(1)

experiment_table = pd.DataFrame({
    "count": experiment_counts,
    "percentage": experiment_pct,
    "description": experiment_counts.index.map(experiment_labels)
})

print("Experiment distribution with descriptions (counts and percentages):")
print(experiment_table)

print("\n" + "=" * 80)

# Summarise parking-rights decisions across the repeated tasks.
give_up_counts = (
    data_formatted[["ResponseId", "give_up_parking"]]
    .assign(give_up_yes=lambda d: d["give_up_parking"] == "Yes")
    .groupby("ResponseId")["give_up_yes"]
    .sum()
)

grouped = data_formatted.groupby('ResponseId')['give_up_parking'] \
                        .agg(lambda x: set(x.dropna()))

only_yes = grouped.apply(lambda x: x == {'Yes'})
only_no = grouped.apply(lambda x: x == {'No'})
both = grouped.apply(lambda x: ('Yes' in x) and ('No' in x))

n_only_yes = only_yes.sum()
n_only_no = only_no.sum()
n_both = both.sum()

print(f"# Count ONLY Yes in give up parking: {n_only_yes}")

print(f"# Count ONLY No in give up parking: {n_only_no}")

print(f"# Count BOTH Yes and No in give up parking: {n_both}")

total_respondents = grouped.shape[0]
print(f"# Total respondents: {total_respondents}")

df_car = data_formatted[
    data_formatted["primary_transport_mode"].isin(["car_asml", "car_pr"])
].copy()

never_car = df_car.groupby('ResponseId')['base_choice'] \
                  .apply(lambda x: (x != 'car').all())

n_never_car = never_car.sum()

total_car_users = never_car.shape[0]

print(f"# Primary car users who NEVER chose car: {n_never_car}")

print(f"# Total primary car users: {total_car_users}")

print(
    "Average number of times respondents gave up parking:",
    give_up_counts.mean().round(2)
)

give_up_counts = (
    data_formatted[["ResponseId", "primary_transport_mode", "give_up_parking"]]
    .assign(give_up_yes=lambda d: d["give_up_parking"] == "Yes")
    .groupby(["ResponseId", "primary_transport_mode"])["give_up_yes"]
    .sum()
    .reset_index()
)

avg_give_up = give_up_counts.groupby(
    give_up_counts["primary_transport_mode"] == "car_asml"
)["give_up_yes"].mean().round(2)

print("Average number of times respondents gave up parking:")
print("- Car (ASML) as primary mode:", avg_give_up.get(True))
print("- Other primary modes:", avg_give_up.get(False))

print("\n" + "=" * 80)

# Describe the compensation selected after giving up parking rights.
reward_counts = (
    data_formatted[["ResponseId", "reward_choice_if_gave_up_parking"]]
    .drop_duplicates("ResponseId")["reward_choice_if_gave_up_parking"]
    .dropna()
    .value_counts()
)

reward_pct = (reward_counts / reward_counts.sum() * 100).round(1)

reward_table = pd.DataFrame({
    "count": reward_counts,
    "percentage": reward_pct
})

print("Reward choice if respondents gave up parking (counts and percentages, NaNs excluded):")
print(reward_table)

print("\n" + "=" * 80)

# This internal output contains ResponseIds and free-text answers.
# Keep the generated document outside the public repository.
respondent_summary = (
    data_formatted[
        [
            "ResponseId",
            "commute_distance_km",
            "car_available",
            "bike_available",
            "vanpool_available",
            "primary_transport_mode",
            "reason_not_giving_up_parking",
            "additional_comments",
            "gave_up_parking_any"
        ]
    ]
    .drop_duplicates("ResponseId")
    .loc[
        lambda d: ~(
            d["reason_not_giving_up_parking"].isna()
            & d["additional_comments"].isna()
        )
    ]
)

doc = Document()
doc.add_heading("Survey Open Answers – Respondent Overview", level=1)

for _, row in respondent_summary.iterrows():

    doc.add_heading(f"Respondent {row['ResponseId']}", level=2)

    doc.add_paragraph(
        f"Commute distance (km): {row['commute_distance_km']}\n"
        f"Primary transport mode: {row['primary_transport_mode']}\n"
        f"Car available: {row['car_available']}\n"
        f"Bike available: {row['bike_available']}\n"
        f"Vanpool available: {row['vanpool_available']}\n"
        "Did they give up their parking in at least 1/8 scenarios?: "
        f"{row['gave_up_parking_any']}\n"
    )

    if row["gave_up_parking_any"] == "No":
        doc.add_paragraph("Reason for not giving up parking:", style="List Bullet")
        doc.add_paragraph(
            row["reason_not_giving_up_parking"]
            if pd.notna(row["reason_not_giving_up_parking"])
            else "—",
            style="Normal"
        )

    doc.add_paragraph("Additional comments:", style="List Bullet")
    doc.add_paragraph(
        row["additional_comments"] if pd.notna(row["additional_comments"]) else "—",
        style="Normal"
    )

    doc.add_page_break()

doc.save(r"data\Respondent Open Answers.docx")

print("Word document created: respondent_open_answers.docx")

print("\n" + "=" * 80)

print("\n" + "=" * 80)

# Group commute distances into the bands used for reporting.
distance_band_df = (
    data_formatted[["ResponseId", "commute_distance_km"]]
    .drop_duplicates("ResponseId")
    .copy()
)

bins = [0, 5, 10, 20, 30, 50, np.inf]
labels = [
    "<5 km",
    "5-10 km",
    "10-20 km",
    "20-30 km",
    "30-50 km",
    "50+ km"
]

distance_band_df["distance_band"] = pd.cut(
    distance_band_df["commute_distance_km"],
    bins=bins,
    labels=labels,
    right=False
)

distance_counts = distance_band_df["distance_band"].value_counts().sort_index()

distance_pct = (
    distance_band_df["distance_band"]
    .value_counts(normalize=True)
    .sort_index()
    .mul(100)
    .round(1)
)

distance_table = pd.DataFrame({
    "count": distance_counts,
    "percentage": distance_pct
})

print("Commute distance distribution:")
print(distance_table)

print("\n# Percentage of employees per commute distance band:")
for band, pct in distance_pct.items():
    print(f"{band}: {pct:.1f}%")
